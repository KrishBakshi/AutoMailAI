from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import zipfile


def write_email(name, company, ai_email):
    output_folder = "./output"
    # Create a new Word document
    doc = Document()

    # Add AI-generated text to the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(ai_email)

    # Formatting the text
    run.font.name = "Trebuchet MS"
    run.font.size = Pt(10)

    # # Justify the paragraph
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Set vertical spacing (before and after paragraph)
    paragraph.paragraph_format.space_before = Pt(10)  # 10pt space before paragraph
    paragraph.paragraph_format.space_after = Pt(10)  # 10pt space after paragraph
    paragraph.paragraph_format.line_spacing = 1.0  # 1.5 line spacing

    # Save the document
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    file_path = os.path.join(output_folder, f"{name}, {company}.docx")
    doc.save(file_path)

    return file_path  # Return the file path for Gradio to use

def return_bulk_file():
    output_folder = "./output"
    # Create a ZIP file of the output folder
    zip_file_path = "./output.zip"
    with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(output_folder):
            for file in files:
                file_full_path = os.path.join(root, file)
                zipf.write(file_full_path, os.path.relpath(file_full_path, output_folder))

    return zip_file_path  # Return the ZIP file path for Gradio to use

# Define the system instruction

sys_instruct_gen = """
You are an AI assistant tasked with drafting a professional and personalized email to a hiring manager 
expressing interest in a job opportunity at their company. 
You need to adjust accordingly with grammar and spot possible areas where the info in the resume and that of the hiring manager/recruiter are relavent.
Make the email such that it piques the interest of the hiring manger, such as a certain skill set they need or can be of great value to them.
The email should adhere to the following template and should generate in 260 - 280 words consistently, with placeholders to be filled accordingly:

Subject: Seeking Opportunities to Contribute at {Company Name}

Dear {Hiring Manager's Name},

I hope this email finds you well. My name is [Extract Full Name from Resume], and I am writing to express my keen interest in exploring career opportunities at {Company Name}. Your organization's work in {Industry/Field} has greatly impressed me, particularly your contributions to {Specific Projects, Innovations, or Company Achievements}.

With a background in [Extract Academic Background: degree, university, and GPA], I have developed strong skills in [Extract Relevant Technical Skills] and gained hands-on experience through projects such as [Extract Relevant Projects from Resume]. My expertise in {Specific Skills or Tools Relevant to the Job} aligns well with the work being done at your company, and I am eager to bring my knowledge and passion to your team.

Additionally, my certification in Practical AI with Python showcases my commitment to continuous learning in AI and data-driven solutions. Furthermore, my proficiency in Japanese, certified by the JLPT N3, allows me to collaborate effectively in diverse and international work environments.

I am particularly interested in {Specific Roles, Teams, or Projects at the Company} and believe that my skills in {Extract Skills from Resume} would allow me to contribute meaningfully to your organization's goals. I have attached my resume for your review and would welcome the opportunity to discuss how my experience and expertise align with your company’s needs.

Thank you for your time and consideration. I look forward to the possibility of joining {Company Name} and contributing to its continued success.

Best regards,
[Extract Full Name from Resume]

Resume:
{}

"""

sys_instruct_jap = """
You are an AI assistant tasked with drafting a professional and personalized email to a hiring manager 
expressing interest in a job opportunity at their company. 
You need to adjust accordingly with grammar and spot possible areas where the info in the resume and that of the hiring manager/recruiter are relavent.
Make the email such that it piques the interest of the hiring manger, such as a certain skill set they need or can be of great value to them.
Always include japanese language point present in the template.
The email should adhere to the following template and should generate in 240 - 260 words consistently, with placeholders to be filled accordingly:

Subject: Seeking Opportunities to Contribute at {Company Name}

Dear {Hiring Manager's Name},

I hope this email finds you well. My name is [Extract Full Name from Resume], and I am writing to express my keen interest in exploring career opportunities at {Company Name}. Your organization's work in {Industry/Field} has greatly impressed me, particularly your contributions to {Specific Projects, Innovations, or Company Achievements}.

With a background in [Extract Academic Background: degree, university, and GPA], I have developed strong skills in [Extract Relevant Technical Skills] and gained hands-on experience through projects such as [Extract Relevant Projects from Resume]. 

My expertise in {Specific Skills or Tools Relevant to the Job} aligns well with the work being done at your company, and I am eager to bring my knowledge and passion to your team.

Furthermore, my proficiency in Japanese, certified by the JLPT N3, allows me to collaborate effectively in diverse and international work environments.

I am particularly interested in {Specific Roles, Teams, or Projects at the Company} and believe that my skills in {Extract Skills from Resume} would allow me to contribute meaningfully to your organization's goals. I have attached my resume for your review and would welcome the opportunity to discuss how my experience and expertise align with your company’s needs.

Thank you for your time and consideration. I look forward to the possibility of joining {Company Name} and contributing to its continued success.

Best regards,
[Extract Full Name from Resume]

Resume:
{}

"""

prf_tem = sys_instruct_prof_email = """
You are an AI assistant tasked with drafting a professional and personalized email to a professor 
expressing interest in contributing to their research lab and applying for a research internship or master's program opportunity.

You must extract relevant candidate details from the attached resume and adjust the email for clarity, professionalism, and personalization. 
Align the candidate’s skills, academic background, and experience with the professor’s research themes and achievements.

Always retain this fixed element:
1. Mention of JLPT N3 Japanese proficiency.

The generated email must be **formal, inspiring, and consistently between 240–260 words**, with placeholders as defined below:

---

Subject: Interest in Contributing to Your Research Lab

Dear {Professor's Name},

I hope this message finds you well. My name is [Extract Full Name from Resume], and I am writing to express my sincere interest in contributing to the impactful and innovative research conducted at your laboratory. Your work in {Professor's Research Areas} has greatly inspired me and aligns deeply with my academic background and career aspirations.

Your groundbreaking research in {Specific Projects or Research Topics}, along with your achievements such as {Notable Achievements, Awards, or Publications}, exemplifies your leadership and dedication to advancing {Field or Domain}. I am particularly drawn to your work in {Specific Research Details}, as it closely relates to my experience in [Extract Relevant Projects and Experience from Resume].

I am currently pursuing [Extract Academic Background: degree, university, and GPA]. With a strong foundation in [Extract Technical Skills], I have worked on projects such as [Extract Relevant Projects], which have equipped me with hands-on experience in {Skills/Tools Relevant to Their Research}.

Further, my proficiency in Japanese, certified by the JLPT N3, positions me as a strong candidate to integrate seamlessly into your research environment and collaborate effectively within a Japanese academic setting.

I am particularly interested in exploring a **research internship or assistantship opportunity** in your lab, where I can further develop my skills while contributing meaningfully to your ongoing work. I believe my strengths in {Extract Skills from Resume} and my passion for applying technology to real-world challenges make me a valuable addition to your team.

I have attached my resume for your review and would be honored to discuss my qualifications and potential contributions further.

Thank you for considering my application. I look forward to the opportunity to collaborate and learn under your guidance.

Sincerely,  
[Extract Full Name from Resume]

Resume:  
{}
"""


template = {
    "General": sys_instruct_gen,
    "Japanese specific": sys_instruct_jap
}